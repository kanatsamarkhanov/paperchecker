import streamlit as st
from docx import Document
from docx.oxml.ns import qn
from PIL import Image
import re
import pandas as pd
from io import BytesIO
import datetime

st.set_page_config(page_title="Article Checker / Мақала тексеру", page_icon="📋", layout="wide")

if "lang"  not in st.session_state: st.session_state.lang  = "kz"
if "theme" not in st.session_state: st.session_state.theme = "light"

locales = {
  "ru": {
    "title":"📋 Автоматическая проверка статьи",
    "subtitle":"Вестник ЕНУ им. Л.Н. Гумилева · Химия / География · 2025",
    "btn_theme_dark":"🌙 Тёмная тема","btn_theme_light":"☀️ Светлая тема",
    "upload_title":"📂 Загрузите статью в формате .docx",
    "upload_help":"Шаблон журнала Вестник ЕНУ, серия Химия/География, 2025",
    "analyzing":"Анализируем статью...","res_title":"📊 Результаты проверки",
    "total":"Всего","passed":"✅ Выполнено","warned":"⚠️ Внимание",
    "failed":"❌ Не выпол.","score":"🏆 Соответствие",
    "det_report":"### 📋 Детальный отчёт","img_report":"### 🖼️ Анализ рисунков",
    "img_num":"№","img_pixels":"Пиксели","img_size_mm":"Размер в doc",
    "img_dpi_calc":"DPI (расч.)","img_dpi_emb":"DPI (встр.)","img_dpi_real":"DPI реальный",
    "img_format":"Формат","img_status":"Статус",
    "img_label":"Номер рисунка",
    "img_caption":"Подпись под рисунком",
    "img_ref":"Ссылок в тексте",
    "img_width":"Ширина (см)",
    "img_height":"Высота (см)",
    "img_composite":"Возможно составной рисунок",
    "img_capbold":"Подпись жирная",
    "tbl_label":"Номер таблицы",
    "tbl_caption":"Подпись таблицы",
    "tbl_ref":"Ссылок в тексте",
    "tbl_capabove":"Подпись над таблицей",
    "tbl_headbold":"Заголовки жирные",
    "btn_csv":"⬇️ Скачать CSV","btn_xls":"⬇️ Скачать Excel","btn_docx":"⬇️ Word (DOCX)",
    "btn_csv_fig":"⬇️ CSV (рисунки)","btn_csv_tbl":"⬇️ CSV (таблицы)",
    "req_fix":"### ⚠️ Требует исправления","req":"требование",
    "no_file":"👆 Загрузите .docx файл, чтобы начать проверку",
    "c_title":"Наименование статьи","c_title_req":"Строки 3–4 документа",
    "c_lang":"Язык статьи","c_lang_req":"По названию статьи",
    "c_vol":"Объём статьи","c_vol_req":"≥3500 слов",
    "c_ann_main":"Основная аннотация","c_ann_req":"≤300 слов",
    "c_ann_ru":"Аннотация (рус)","c_ann_kz":"Аннотация (каз)","c_ann_en":"Abstract (англ)",
    "c_req_obl":"Обязательно",
    "c_kw":"Ключевые слова","c_kw_req":"3–10, разделитель «;»",
    "c_mrnti":"Код МРНТИ / IRSTI","c_orcid":"ORCID авторов","c_orcid_req":"Для каждого автора",
    "c_intro":"§1. Введение","c_mm":"§2. Материалы и методы",
    "c_res":"§3. Результаты","c_disc":"§4. Талдау",
    "c_concl":"§5. Заключение",
    "c_supp":"§6. Вспомог. материал","c_contrib":"§7. Вклад авторов",
    "c_authinfo":"§8. Информация об авторе","c_fund":"§9. Финансирование",
    "c_ack":"§10. Благодарности","c_conflict":"§11. Конфликты интересов",
    "c_paper":"Формат бумаги","c_paper_req":"A4 (210x297 мм)",
    "c_margins":"Поля","c_margins_req":"Все поля 20 мм",
    "c_font":"Шрифт и кегль","c_font_req":"Times New Roman, 12 pt",
    "c_tables":"Таблицы","c_tables_req":"Должны быть в тексте",
    "c_images":"Рисунки","c_images_req":"600 DPI, TIFF/JPEG/PNG",
    "c_multi_ann":"Многоязычные аннотации","c_multi_ann_req":"Ещё 2 аннотации на других языках",
    "c_ref_apa":"Литература: стиль APA",
    "c_ref_apa_req":"Формат Author, A. A. (год).",
    "c_ref_age":"Литература: давность источников",
    "c_ref_age_req_prefix":"Годы ≥ ",
    "img_see_table":"см. таблицу ниже",
    "found":"Найдено","not_found":"Отсутствует","words":"слов",
    "f_author":"Канат Самарханов / Kanat Samarkhanov","f_license":"Лицензия",
    "f_univ":"ЕНУ им. Л.Н. Гумилева — Кафедра физической и экономической географии",
  },
  "kz": {
    "title":"📋 Мақаланы автоматты түрде тексеру",
    "subtitle":"Л.Н. Гумилев атындағы ЕҰУ Хабаршысы · Химия / География · 2025 үлгісі",
    "btn_theme_dark":"🌙 Түнгі режим","btn_theme_light":"☀️ Күндізгі режим",
    "upload_title":"📂 .docx форматындағы мақаланы жүктеңіз",
    "upload_help":"Л.Н. Гумилев атындағы ЕҰУ Хабаршысы, Химия/География сериясы, 2025 үлгісі",
    "analyzing":"Мақала талдануда...","res_title":"📊 Тексеру нәтижелері",
    "total":"Барлығы","passed":"✅ Орындалды","warned":"⚠️ Назар аударыңыз",
    "failed":"❌ Орындалмады","score":"🏆 Сәйкестік",
    "det_report":"### 📋 Толық есеп","img_report":"### 🖼️ Суреттерді талдау",
    "img_num":"№","img_pixels":"Пикселдер","img_size_mm":"Doc өлшемі",
    "img_dpi_calc":"DPI (есепт.)","img_dpi_emb":"DPI (енгіз.)","img_dpi_real":"Нақты DPI",
    "img_format":"Формат","img_status":"Статус",
    "img_label":"Сурет нөмірі",
    "img_caption":"Сурет астындағы мәтін",
    "img_ref":"Мәтінде сілтемелер саны",
    "img_width":"Ені (см)",
    "img_height":"Биіктігі (см)",
    "img_composite":"Құрама сурет болуы мүмкін",
    "img_capbold":"Қалың қаріппен бе",
    "tbl_label":"Кесте нөмірі",
    "tbl_caption":"Кесте атауы",
    "tbl_ref":"Мәтінде сілтемелер саны",
    "tbl_capabove":"Жазуы кестенің үстінде",
    "tbl_headbold":"Баған атаулары қалың",
    "btn_csv":"⬇️ CSV жүктеу","btn_xls":"⬇️ Excel жүктеу","btn_docx":"⬇️ Word (DOCX)",
    "btn_csv_fig":"⬇️ CSV (суреттер)","btn_csv_tbl":"⬇️ CSV (кестелер)",
    "req_fix":"### ⚠️ Түзетуді қажет етеді","req":"талап",
    "no_file":"👆 Тексеруді бастау үшін .docx файлын жүктеңіз",
    "c_title":"Мақаланың атауы","c_title_req":"Құжаттың 3–4 жолдары",
    "c_lang":"Мақала тілі","c_lang_req":"Мақала атауына қарай",
    "c_vol":"Мақала көлемі","c_vol_req":"≥3500 сөз",
    "c_ann_main":"Негізгі аңдатпа","c_ann_req":"≤300 сөз",
    "c_ann_ru":"Аңдатпа (орыс)","c_ann_kz":"Аңдатпа (қаз)","c_ann_en":"Abstract (ағылш)",
    "c_req_obl":"Міндетті түрде",
    "c_kw":"Түйінді сөздер","c_kw_req":"3–10, бөлгіш «;»",
    "c_mrnti":"МРНТИ / IRSTI коды","c_orcid":"Авторлардың ORCID","c_orcid_req":"Әр автор үшін",
    "c_intro":"§1. Кіріспе","c_mm":"§2. Материалдар мен әдістер",
    "c_res":"§3. Нәтижелер","c_disc":"§4. Талдау",
    "c_concl":"§5. Қорытынды",
    "c_supp":"§6. Қосымша материалдар","c_contrib":"§7. Авторлардың үлесі",
    "c_authinfo":"§8. Автор туралы ақпарат","c_fund":"§9. Қаржыландыру",
    "c_ack":"§10. Алғыстар","c_conflict":"§11. Мүдделер қақтығысы",
    "c_paper":"Қағаз форматы","c_paper_req":"A4 (210x297 мм)",
    "c_margins":"Жақтаулар","c_margins_req":"Барлық жақтаулар 20 мм",
    "c_font":"Шрифт және кегль","c_font_req":"Times New Roman, 12 pt",
    "c_tables":"Кестелер","c_tables_req":"Мәтінде болуы керек",
    "c_images":"Суреттер","c_images_req":"600 DPI, TIFF/JPEG/PNG",
    "c_multi_ann":"Көптілді аңдатпалар","c_multi_ann_req":"Басқа 2 тілде аңдатпа болуы керек",
    "c_ref_apa":"Әдебиет: APA стилі",
    "c_ref_apa_req":"Формат Author, A. A. (жыл).",
    "c_ref_age":"Әдебиет: дереккөздердің жаңалығы",
    "c_ref_age_req_prefix":"Жылдар ≥ ",
    "img_see_table":"төмендегі кестені қараңыз",
    "found":"Табылды","not_found":"Жоқ","words":"сөз",
    "f_author":"Канат Самарханов / Kanat Samarkhanov","f_license":"Лицензия",
    "f_univ":"Л.Н. Гумилев атындағы ЕҰУ — Физикалық және экономикалық география кафедрасы",
  },
  "en": {
    "title":"📋 Article Compliance Checker",
    "subtitle":"L.N. Gumilyov ENU Bulletin · Chemistry / Geography · 2025 Template",
    "btn_theme_dark":"🌙 Dark mode","btn_theme_light":"☀️ Light mode",
    "upload_title":"📂 Upload article in .docx format",
    "upload_help":"L.N. Gumilyov ENU Bulletin, Chemistry/Geography series, 2025 template",
    "analyzing":"Analysing article...","res_title":"📊 Compliance Results",
    "total":"Total","passed":"✅ Passed","warned":"⚠️ Warning",
    "failed":"❌ Failed","score":"🏆 Score",
    "det_report":"### 📋 Detailed Report","img_report":"### 🖼️ Figure Analysis",
    "img_num":"No.","img_pixels":"Pixels","img_size_mm":"Size in doc",
    "img_dpi_calc":"DPI (calc.)","img_dpi_emb":"DPI (emb.)","img_dpi_real":"Real DPI",
    "img_format":"Format","img_status":"Status",
    "img_label":"Figure number",
    "img_caption":"Figure caption",
    "img_ref":"References in text",
    "img_width":"Width (cm)",
    "img_height":"Height (cm)",
    "img_composite":"Possibly composite figure",
    "img_capbold":"Caption in bold",
    "tbl_label":"Table number",
    "tbl_caption":"Table caption",
    "tbl_ref":"References in text",
    "tbl_capabove":"Caption above table",
    "tbl_headbold":"Header row in bold",
    "btn_csv":"⬇️ Download CSV","btn_xls":"⬇️ Download Excel","btn_docx":"⬇️ Word (DOCX)",
    "btn_csv_fig":"⬇️ CSV (figures)","btn_csv_tbl":"⬇️ CSV (tables)",
    "req_fix":"### ⚠️ Requires Correction","req":"requirement",
    "no_file":"👆 Upload a .docx file to start checking",
    "c_title":"Article title","c_title_req":"Lines 3–4 of document",
    "c_lang":"Article language","c_lang_req":"By article title",
    "c_vol":"Article volume","c_vol_req":"≥3500 words",
    "c_ann_main":"Main abstract","c_ann_req":"≤300 words",
    "c_ann_ru":"Abstract (rus)","c_ann_kz":"Abstract (kaz)","c_ann_en":"Abstract (eng)",
    "c_req_obl":"Required",
    "c_kw":"Keywords","c_kw_req":"3–10, separator ;",
    "c_mrnti":"IRSTI / МРНТИ code","c_orcid":"Author ORCIDs","c_orcid_req":"Per each author",
    "c_intro":"§1. Introduction","c_mm":"§2. Materials & Methods",
    "c_res":"§3. Results","c_disc":"§4. Discussion",
    "c_concl":"§5. Conclusion",
    "c_supp":"§6. Supplementary Material","c_contrib":"§7. Author Contributions",
    "c_authinfo":"§8. Author Information","c_fund":"§9. Funding",
    "c_ack":"§10. Acknowledgements","c_conflict":"§11. Conflicts of Interest",
    "c_paper":"Paper format","c_paper_req":"A4 (210x297 mm)",
    "c_margins":"Margins","c_margins_req":"All margins 20 mm",
    "c_font":"Font & size","c_font_req":"Times New Roman, 12 pt",
    "c_tables":"Tables","c_tables_req":"Must be in text",
    "c_images":"Figures","c_images_req":"600 DPI, TIFF/JPEG/PNG",
    "c_multi_ann":"Multilingual abstracts","c_multi_ann_req":"2 more abstracts in other languages",
    "c_ref_apa":"References: APA style",
    "c_ref_apa_req":"Format Author, A. A. (year).",
    "c_ref_age":"References: recency of sources",
    "c_ref_age_req_prefix":"Years ≥ ",
    "img_see_table":"see table below",
    "found":"Found","not_found":"Not found","words":"words",
    "f_author":"Kanat Samarkhanov","f_license":"License",
    "f_univ":"L.N. Gumilyov ENU — Department of Physical and Economic Geography",
  },
}
l = locales[st.session_state.lang]

dark_css = (
    "<style>"
    "html,body,[class*='css'],.stApp{background-color:#0d1b2e !important;color:#c9d8ee !important;"
    "font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Helvetica,Arial,sans-serif !important;}"
    "h1,h2,h3,h4,h5,h6,[data-testid='stMarkdownContainer'] h1,[data-testid='stMarkdownContainer'] h2,"
    "[data-testid='stMarkdownContainer'] h3{color:#e2edf7 !important;font-weight:600 !important;}"
    "p,span,label,div,li,[data-testid='stMarkdownContainer'] p,"
    "[data-testid='stCaptionContainer'],.stCaption{color:#c9d8ee !important;}"
    "[data-testid='block-container'],[data-testid='stVerticalBlock'],"
    "section[data-testid='stSidebar']{background-color:#0d1b2e !important;}"
    "[data-testid='stMetric']{background:#0f2340 !important;border:1px solid #1e3a5f !important;"
    "border-radius:6px !important;padding:12px 16px !important;}"
    "[data-testid='stMetricValue']{color:#e2edf7 !important;}"
    "[data-testid='stMetricLabel']{color:#7b96b8 !important;}"
    ".stButton>button{background-color:#0f2340 !important;color:#c9d8ee !important;"
    "border:1px solid #1e3a5f !important;border-radius:6px !important;}"
    ".stButton>button:hover{background-color:#1e3a5f !important;color:#e2edf7 !important;}"
    "[data-testid='stDownloadButton']>button{background-color:#238636 !important;color:#fff !important;"
    "border:1px solid #2ea043 !important;border-radius:6px !important;}"
    "[data-testid='stDownloadButton']>button:hover{background-color:#2ea043 !important;}"
    "[data-testid='stFileUploader']{background-color:#0f2340 !important;border-radius:8px !important;}"
    "[data-testid='stFileUploaderDropzone']{background-color:#0f2340 !important;"
    "border:2px dashed #1e3a5f !important;border-radius:8px !important;padding:24px 16px !important;}"
    "[data-testid='stFileUploaderDropzone']:hover{background-color:#112850 !important;border-color:#2f5f9e !important;}"
    "[data-testid='stFileUploader'] *,[data-testid='stFileUploaderDropzone'] *{color:#c9d8ee !important;}"
    "[data-testid='stFileUploaderDropzone'] button{background-color:#1e3a5f !important;"
    "color:#c9d8ee !important;border:1px solid #2f5f9e !important;border-radius:6px !important;"
    "padding:5px 16px !important;font-size:13px !important;font-weight:500 !important;}"
    "[data-testid='stFileUploaderDropzone'] button:hover{background-color:#2f5f9e !important;"
    "border-color:#58a6ff !important;color:#e2edf7 !important;}"
    "[data-testid='stFileUploaderFile']{background-color:#112240 !important;"
    "border:1px solid #1e3a5f !important;border-radius:6px !important;}"
    "[data-testid='stFileUploaderDeleteBtn'] button{color:#7b96b8 !important;}"
    "[data-testid='stFileUploaderDeleteBtn'] button:hover{color:#f85149 !important;}"
    "[data-testid='stDataFrame'],.stDataFrame iframe{border:1px solid #1e3a5f !important;"
    "border-radius:8px !important;overflow:hidden !important;"
    "box-shadow:0 2px 8px rgba(0,0,0,0.4) !important;}"
    "[data-testid='stAlert']{background-color:#0f2340 !important;border:1px solid #1f6feb !important;"
    "color:#c9d8ee !important;border-radius:6px !important;}"
    ".stSpinner>div{color:#c9d8ee !important;}"
    "hr{border-color:#1e3a5f !important;}"
    "input,textarea,select{background-color:#0f2340 !important;color:#c9d8ee !important;"
    "border:1px solid #1e3a5f !important;}"
    "[data-testid='stSelectbox']>div>div{background-color:#0f2340 !important;"
    "border:1px solid #1e3a5f !important;border-radius:6px !important;color:#c9d8ee !important;}"
    "</style>"
)

light_css = (
    "<style>"
    "[data-testid='stMetric']{background:#fff;padding:12px;border-radius:10px;box-shadow:0 2px 6px rgba(0,0,0,.08);}"
    "h1,h2,h3{color:#1a3a5c;}"
    "[data-testid='stDownloadButton']>button{background-color:#2ea043;color:#fff;border-radius:6px;}"
    "[data-testid='stDataFrame'],.stDataFrame iframe{border:1px solid #d0d7de;"
    "border-radius:8px;overflow:hidden;box-shadow:0 1px 4px rgba(0,0,0,0.08);}"
    "</style>"
)

st.markdown(dark_css if st.session_state.theme == "dark" else light_css, unsafe_allow_html=True)

_DB_CARD="#0f2340"; _DB_HEAD="#e2edf7"; _DB_MUTED="#7b96b8"

hc1, hc2, hc3 = st.columns([6, 1.8, 1.8])
with hc1:
    st.title(l["title"])
    st.caption(l["subtitle"])
with hc2:
    _lang_labels = {"kz": "🇰🇿 Қазақша", "ru": "🇷🇺 Русский", "en": "🇬🇧 English"}
    _lang_keys   = list(_lang_labels.keys())
    _sel = st.selectbox(
        "lang", _lang_keys,
        index=_lang_keys.index(st.session_state.lang),
        format_func=lambda x: _lang_labels[x],
        label_visibility="collapsed",
    )
    if _sel != st.session_state.lang:
        st.session_state.lang = _sel
        st.rerun()
with hc3:
    _tbtn = l["btn_theme_light"] if st.session_state.theme == "dark" else l["btn_theme_dark"]
    if st.button(_tbtn, use_container_width=True):
        st.session_state.theme = "light" if st.session_state.theme == "dark" else "dark"
        st.rerun()
st.markdown("---")

_KAZ_CHARS = set("қңөұүіәғҚҢӨҰҮІӘҒ")

def detect_lang_from_text(text):
    latin = sum(1 for c in text if c.isalpha() and c.isascii())
    cyr   = sum(1 for c in text if "\u0400" <= c <= "\u04FF")
    kaz   = sum(1 for c in text if c in _KAZ_CHARS)
    if latin > cyr: return "en"
    if kaz >= 1:    return "kz"
    return "ru"

def extract_title_and_lang(doc):
    non_empty  = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    candidates = non_empty[2:4] if len(non_empty) >= 3 else non_empty
    title = max(candidates, key=len) if candidates else ""
    return title, detect_lang_from_text(title)

def title_for_filename(title):
    if not title: return "report"
    clean = re.sub(r"[^\w\s-]", "", title, flags=re.UNICODE)
    return clean.strip().replace(" ", "_")[:40] or "report"

_CONFLICT_RE = re.compile(
    r"конфликт\w*\s+интерес|conflict\w*\s+of\s+interest|мүдделер\s+қақтығыс",
    re.IGNORECASE,
)

def has_conflict_section(doc, full_text):
    if _CONFLICT_RE.search(full_text): return True
    for par in doc.paragraphs:
        if _CONFLICT_RE.search(par.text): return True
    return False

# ── Abstract extraction ───────────────────────────────────────────────────
_ANN_END = r"(?=ключевые\s+слова|keywords|түйінді\s+сөздер|түйін\s+сөздер|введение|кіріспе|introduction|\Z)"

_ANN_PATTERNS = {
    "ru": re.compile(
        r"аннотация\s*[:\—\-]?\s*(.{30,}?)" + _ANN_END,
        re.IGNORECASE | re.DOTALL,
    ),
    "kz": re.compile(
        r"аңдатпа\s*[:\—\-]?\s*(.{30,}?)" + _ANN_END,
        re.IGNORECASE | re.DOTALL,
    ),
    "en": re.compile(
        r"\babstract\b\s*[:\—\-]?\s*(.{30,}?)" + _ANN_END,
        re.IGNORECASE | re.DOTALL,
    ),
}

def extract_abstract(full_text, lang, region=None):
    txt = full_text
    m = _ANN_PATTERNS[lang].search(txt)
    return m.group(1).strip() if m else None

# ── Image & table analysis ────────────────────────────────────────────────
_ALLOWED_FORMATS = {"TIFF", "JPEG", "PNG"}
_MIN_DPI = 600

_CAPTION_FIG_RE = re.compile(
    r"^\s*(Figure|Сурет|Рисунок)\s+(\d+)\s*\.", re.IGNORECASE | re.MULTILINE
)
_CAPTION_TBL_RE = re.compile(
    r"^\s*(Table|Кесте|Таблица)\s+(\d+)\s*\.", re.IGNORECASE | re.MULTILINE
)

_REF_FIG_RE = re.compile(
    r"\b(Figure|Fig\.?|Сурет|Рисунок)\s+(\d+)\b", re.IGNORECASE
)
_REF_TBL_RE = re.compile(
    r"\b(Table|Таблица|Табл\.?|Кесте)\s+(\d+)\b", re.IGNORECASE
)

def _yes_no(val, lang_code):
    if lang_code == "ru":
        return "Да" if val else "Нет"
    if lang_code == "kz":
        return "Иә" if val else "Жоқ"
    return "Yes" if val else "No"

def analyse_figures_and_tables(doc, full_text, l):
    EMU = 914400
    img_rows = []
    tables_rows = []

    fig_refs = {}
    for m in _REF_FIG_RE.finditer(full_text):
        num = int(m.group(2))
        fig_refs[num] = fig_refs.get(num, 0) + 1

    tbl_refs = {}
    for m in _REF_TBL_RE.finditer(full_text):
        num = int(m.group(2))
        tbl_refs[num] = tbl_refs.get(num, 0) + 1

    paras = list(doc.paragraphs)
    lang_code = st.session_state.lang

    # рисунки
    for idx, shape in enumerate(doc.inline_shapes):
        w_in = (shape.width or 0) / EMU
        h_in = (shape.height or 0) / EMU
        w_cm = round(w_in * 2.54, 2) if w_in > 0 else None
        h_cm = round(h_in * 2.54, 2) if h_in > 0 else None
        size_mm = f"{round(w_in*25.4)}x{round(h_in*25.4)} mm" if w_in > 0 else "-"

        par_idx = None
        for pi, p in enumerate(paras):
            if shape._inline in p._p:
                par_idx = pi
                break

        caption = ""
        label_num = None
        cap_bold = False
        is_composite = False

        if par_idx is not None and par_idx + 1 < len(paras):
            cap_par = paras[par_idx+1]
            cap_text = cap_par.text.strip()
            caption = cap_text
            m_cap = _CAPTION_FIG_RE.search(cap_text)
            if m_cap:
                label_num = int(m_cap.group(2))
            for run in cap_par.runs:
                if run.text.strip() and run.bold:
                    cap_bold = True
                    break

        if label_num is None:
            label_num = idx + 1
        if not caption:
            is_composite = True

        try:
            pic  = shape._inline.graphic.graphicData.pic
            rId  = pic.blipFill.blip.get(qn("r:embed"))
            blob = doc.part.related_parts[rId].blob
            img  = Image.open(BytesIO(blob))
            px_w, px_h = img.size
            dw = round(px_w / w_in) if w_in > 0 else None
            dh = round(px_h / h_in) if h_in > 0 else None
            dpi_calc = f"{dw}x{dh}" if dw else "-"
            emb = img.info.get("dpi")
            dpi_emb = f"{round(emb[0])}x{round(emb[1])}" if emb else "-"
            fmt = (img.format or "?").upper()
            fmt_ok = fmt in _ALLOWED_FORMATS
            dpi_ok = isinstance(dw, int) and dw >= _MIN_DPI
            if fmt_ok and dpi_ok:
                status = "✅"
            elif not fmt_ok and not dpi_ok:
                status = "❌"
            else:
                status = "⚠️"
            real_dpi = dw if dw else (round(emb[0]) if emb else None)
            real_dpi_str = str(real_dpi) if real_dpi else "-"
        except Exception:
            px_w = px_h = 0
            dpi_calc = dpi_emb = real_dpi_str = fmt = "-"
            status = "⚠️"

        ref_count = fig_refs.get(label_num, 0)

        img_rows.append({
            l["img_num"]: idx + 1,
            l["img_label"]: f"Figure {label_num}",
            l["img_pixels"]: f"{px_w}x{px_h}" if px_w and px_h else "-",
            l["img_size_mm"]: size_mm,
            l["img_width"]:  f"{w_cm}" if w_cm else "-",
            l["img_height"]: f"{h_cm}" if h_cm else "-",
            l["img_dpi_calc"]: dpi_calc,
            l["img_dpi_emb"]: dpi_emb,
            l["img_dpi_real"]: real_dpi_str,
            l["img_format"]: fmt,
            l["img_caption"]: caption or l["not_found"],
            l["img_ref"]: ref_count,
            l["img_composite"]: _yes_no(is_composite, lang_code),
            l["img_capbold"]: _yes_no(cap_bold, lang_code),
            l["img_status"]: status,
        })

    # таблицы
    for t_idx, table in enumerate(doc.tables, start=1):
        cap = ""
        tbl_par_idx = None
        for pi, p in enumerate(paras):
            if table._tbl in p._p:
                tbl_par_idx = pi
                break
        caption_above = False
        header_bold_ok = True

        if tbl_par_idx is not None and tbl_par_idx > 0:
            tx = paras[tbl_par_idx-1].text.strip()
            if tx:
                cap = tx
                caption_above = True

        m_cap = _CAPTION_TBL_RE.search(cap) if cap else None
        tnum = int(m_cap.group(2)) if m_cap else t_idx
        ref_count = tbl_refs.get(tnum, 0)

        if table.rows:
            header_row = table.rows[0]
            for cell in header_row.cells:
                cell_text = cell.text.strip()
                if not cell_text:
                    continue
                any_bold = False
                for p in cell.paragraphs:
                    for r in p.runs:
                        if r.text.strip() and r.bold:
                            any_bold = True
                            break
                    if any_bold:
                        break
                if not any_bold:
                    header_bold_ok = False
                    break

        tables_rows.append({
            l["tbl_label"]: f"Table {tnum}",
            l["tbl_caption"]: cap or l["not_found"],
            l["tbl_ref"]: ref_count,
            l["tbl_capabove"]: _yes_no(caption_above, lang_code),
            l["tbl_headbold"]: _yes_no(header_bold_ok, lang_code),
        })

    return img_rows, tables_rows

def detect_author_count(doc, orcid_count):
    if orcid_count >= 1: return orcid_count
    non_empty = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for para in non_empty[3:10]:
        parts = [x.strip() for x in para.split(",") if x.strip()]
        if 2 <= len(parts) <= 8 and all(len(p) < 60 for p in parts):
            return len(parts)
    return 1

_ALL_LANGS   = ["ru", "kz", "en"]
_LANG_LABELS = {"ru": "Русский", "kz": "Қазақша", "en": "English"}
_ANN_KEYS    = {"ru": "c_ann_ru", "kz": "c_ann_kz", "en": "c_ann_en"}

def check_article(doc, l):
    full_text  = "\n".join(p.text for p in doc.paragraphs)
    word_count = len(full_text.split())
    text_low   = full_text.lower()
    results    = []
    title, main_lang = extract_title_and_lang(doc)

    def add(num, crit, req, found, status):
        results.append({"№": num, "Критерий": crit, "Требование": req,
                        "Найдено": found, "Статус": status})

    add(0, l["c_title"], l["c_title_req"], title or l["not_found"], "✅" if title else "⚠️")
    add(1, l["c_lang"],  l["c_lang_req"],  _LANG_LABELS.get(main_lang, main_lang), "✅")
    add(2, l["c_vol"],   l["c_vol_req"],   f"{word_count} {l['words']}",
        "✅" if word_count >= 3500 else "⚠️")

    other_langs = [lg for lg in _ALL_LANGS if lg != main_lang]
    main_ann_text = extract_abstract(full_text, main_lang, region=None)
    main_label    = f"{l['c_ann_main']} ({_LANG_LABELS[main_lang]})"
    if main_ann_text:
        aw = len(main_ann_text.split())
        add(3, main_label, l["c_ann_req"],
            f"{aw} {l['words']}", "✅" if aw <= 300 else "❌")
    else:
        add(3, main_label, l["c_ann_req"], l["not_found"], "⚠️")

    has_other = {}
    for num, olang in zip([4, 5], other_langs):
        ann_text = extract_abstract(full_text, olang, region=None)
        has_other[olang] = ann_text is not None
        add(num, l[_ANN_KEYS[olang]], l["c_req_obl"],
            l["found"] if has_other[olang] else l["not_found"],
            "✅" if has_other[olang] else "❌")

    kw = re.search(r"(ключевые слова|keywords|түйінді сөздер|түйін сөздер)[:\s]+(.+?)(\n|$)",
                   full_text, re.IGNORECASE)
    if kw:
        kw_list = [k.strip() for k in kw.group(2).split(";") if k.strip()]
        add(6, l["c_kw"], l["c_kw_req"], str(len(kw_list)),
            "✅" if 3 <= len(kw_list) <= 10 else "❌")
    else:
        add(6, l["c_kw"], l["c_kw_req"], l["not_found"], "⚠️")

    mrnt = bool(re.search(r"МРНТИ|IRSTI|\d{2}\.\d{2}\.\d{2}", full_text))
    add(7, l["c_mrnti"], l["c_req_obl"], l["found"] if mrnt else l["not_found"],
        "✅" if mrnt else "⚠️")

    orcid = len(re.findall(r"orcid\.org/\d{4}-\d{4}-\d{4}-\d{4}", full_text, re.IGNORECASE))
    add(8, l["c_orcid"], l["c_orcid_req"], f"{orcid} ORCID", "✅" if orcid >= 1 else "⚠️")
    author_count = detect_author_count(doc, orcid)

    for num, name, keys in [
        (9,  l["c_intro"], ["введение","кіріспе","introduction"]),
        (10, l["c_mm"],    ["материалы и методы","материалдар мен әдістер",
                            "materials and methods","материал","әдістер"]),
        (11, l["c_res"],   ["результаты","нәтижелер","results"]),
        (12, l["c_disc"],  ["обсуждение","талқылау","талдау","discussion"]),
        (13, l["c_concl"], ["заключение","қорытынды","conclusion"]),
    ]:
        f = any(k in text_low for k in keys)
        add(num, name, l["c_req_obl"], l["found"] if f else l["not_found"],
            "✅" if f else "❌")

    for num, key, kws, st_ok in [
        (14, l["c_supp"],    ["вспомогательный материал","қосымша материалдар",
                              "supplementary materials"], "⚠️"),
        (16, l["c_authinfo"],["информация об авторе","author information",
                              "автор туралы ақпарат"], "⚠️"),
        (17, l["c_fund"],    ["финансирование","funding","қаржыландыру"], "❌"),
        (18, l["c_ack"],     ["благодарност","acknowledgements","acknowledgments","алғыстар"], "⚠️"),
    ]:
        f = any(k in text_low for k in kws)
        add(num, key, l["c_req_obl"], l["found"] if f else l["not_found"],
            "✅" if f else st_ok)

    contrib_kws   = ["вклад авторов","author contributions","авторлардың үлесі","авторлық үлестер"]
    contrib_found = any(k in text_low for k in contrib_kws)
    if author_count > 1:
        contrib_req    = f"CRediT ({author_count})"
        contrib_status = "✅" if contrib_found else "❌"
    else:
        contrib_req    = "CRediT (1 — opt.)"
        contrib_status = "✅" if contrib_found else "⚠️"
    add(15, l["c_contrib"], contrib_req,
        l["found"] if contrib_found else l["not_found"], contrib_status)

    conflict = has_conflict_section(doc, full_text)
    add(19, l["c_conflict"], l["c_req_obl"],
        l["found"] if conflict else l["not_found"], "✅" if conflict else "❌")

    refs_match = None
    for pat in [r"список литературы", r"references", r"әдебиет(тер)? тізімі"]:
        m2 = re.search(pat, text_low)
        if m2: refs_match = m2; break

    refs_text = ""
    if refs_match:
        refs_text = full_text[refs_match.end():]
        rl = re.findall(r"\n\s*(\[\d+\]|\d+[.)]) ", refs_text)
        if not rl:
            rl = [ln for ln in refs_text.split("\n") if len(ln.strip()) > 40]
        rc = len(rl)
        add(20, "References / Список литературы", "≥10",
            str(rc), "✅" if rc >= 10 else "⚠️")
    else:
        add(20, "References / Список литературы", l["c_req_obl"], l["not_found"], "❌")

    # 26 уже добавлен ниже; сначала бумага/поля/шрифт/таблицы/рисунки/многоязычие
    try:
        sec   = doc.sections[0]
        w_mm  = round(sec.page_width.mm); h_mm = round(sec.page_height.mm)
        is_a4 = (209 <= w_mm <= 211) and (296 <= h_mm <= 298)
        add(21, l["c_paper"], l["c_paper_req"], f"{w_mm}x{h_mm} mm", "✅" if is_a4 else "❌")
        t,b,lf,rg = (round(sec.top_margin.mm), round(sec.bottom_margin.mm),
                     round(sec.left_margin.mm), round(sec.right_margin.mm))
        add(22, l["c_margins"], l["c_margins_req"], f"L:{lf} R:{rg} T:{t} B:{b} mm",
            "✅" if all(x == 20 for x in [t,b,lf,rg]) else "❌")
    except Exception:
        add(21, l["c_paper"],   l["c_paper_req"],  "Error", "⚠️")
        add(22, l["c_margins"], l["c_margins_req"], "Error", "⚠️")

    try:
        fn   = doc.styles["Normal"].font.name or "?"
        fs   = doc.styles["Normal"].font.size.pt if doc.styles["Normal"].font.size else "?"
        ok_f = "Times New Roman" in str(fn) and fs in [11.0, 12.0]
        add(23, l["c_font"], l["c_font_req"], f"{fn}, {fs} pt", "✅" if ok_f else "⚠️")
    except Exception:
        add(23, l["c_font"], l["c_font_req"], "Error", "⚠️")

    tc = len(doc.tables)
    add(24, l["c_tables"], l["c_tables_req"], str(tc), "✅" if tc > 0 else "⚠️")

    ic  = len(doc.inline_shapes)
    msg = f"{ic} — {l['img_see_table']}" if ic > 0 else l["not_found"]
    add(25, l["c_images"], l["c_images_req"], msg, "⚠️" if ic > 0 else "✅")

    ok_multi = all(has_other.values())
    add(26, l["c_multi_ann"], l["c_multi_ann_req"],
        l["found"] if ok_multi else l["not_found"], "✅" if ok_multi else "❌")

    # 27–28: APA и давность источников
    CURRENT_YEAR = datetime.datetime.now().year
    min_year = max_year = None
    yrs_ok = True
    apa_ok = False

    if refs_text:
        years = re.findall(r"\((\d{4})\)", refs_text)
        years = [int(y) for y in years if 1900 <= int(y) <= CURRENT_YEAR]
        if years:
            min_year = min(years)
            max_year = max(years)
            threshold = CURRENT_YEAR - 10
            old_years = [y for y in years if y < threshold]
            yrs_ok = len(old_years) == 0
        if re.search(r"[A-Z][a-z]+,\s*[A-Z]\.", refs_text):
            apa_ok = True

    yrs_req = f"{l['c_ref_age_req_prefix']}{CURRENT_YEAR - 10}"
    yrs_found = f"min={min_year if min_year else '–'}, max={max_year if max_year else '–'}"

    lang_code = st.session_state.lang
    found_txt = {
        "ru": ("Найден шаблон", "Шаблон не найден"),
        "kz": ("Үлгі табылды", "Үлгі табылмады"),
        "en": ("Pattern found", "Pattern not found"),
    }[lang_code]

    add(27, l["c_ref_apa"], l["c_ref_apa_req"],
        found_txt[0] if apa_ok else found_txt[1],
        "✅" if apa_ok else "⚠️")

    add(28, l["c_ref_age"], yrs_req,
        yrs_found, "✅" if yrs_ok else "❌")

    return results, full_text, title, main_lang

def build_docx_report(results, l, total, passed, warned, failed, score,
                      img_rows=None, tables_rows=None):
    buf = BytesIO()
    d = Document()

    d.add_heading(l["res_title"], level=1)
    p = d.add_paragraph()
    p.add_run(f"{l['total']}: {total},  ").bold = True
    p.add_run(
        f"{l['passed']}: {passed},  {l['warned']}: {warned},  "
        f"{l['failed']}: {failed},  {l['score']}: {score}%"
    )
    d.add_paragraph("")

    tbl = d.add_table(rows=1, cols=5)
    for i, h in enumerate(["#", "Criterion", "Requirement", "Found", "Status"]):
        tbl.rows[0].cells[i].text = h
    for r in results:
        row = tbl.add_row().cells
        for i, k in enumerate(["№", "Критерий", "Требование", "Найдено", "Статус"]):
            row[i].text = str(r.get(k, ""))

    if img_rows:
        d.add_page_break()
        d.add_heading(l["img_report"], level=2)
        cols = list(img_rows[0].keys())
        t2 = d.add_table(rows=1, cols=len(cols))
        for i, h in enumerate(cols):
            t2.rows[0].cells[i].text = h
        for row_data in img_rows:
            row = t2.add_row().cells
            for i, key in enumerate(cols):
                row[i].text = str(row_data.get(key, ""))

    if tables_rows:
        d.add_page_break()
        d.add_heading("Tables / Таблицы", level=2)
        cols_t = list(tables_rows[0].keys())
        t3 = d.add_table(rows=1, cols=len(cols_t))
        for i, h in enumerate(cols_t):
            t3.rows[0].cells[i].text = h
        for row_data in tables_rows:
            row = t3.add_row().cells
            for i, key in enumerate(cols_t):
                row[i].text = str(row_data.get(key, ""))

    d.save(buf)
    buf.seek(0)
    return buf.getvalue()

_ST   = {"✅": "background-color:#dafbe1;color:#1a7f37;font-weight:500",
         "⚠️": "background-color:#fff8c5;color:#7d4e00;font-weight:500",
         "❌": "background-color:#ffebe9;color:#cf222e;font-weight:500"}
_BASE = "background-color:#f6f8fa;color:#1f2328"

uploaded_file = st.file_uploader(l["upload_title"], type=["docx"], help=l["upload_help"])

if uploaded_file:
    with st.spinner(l["analyzing"]):
        doc      = Document(uploaded_file)
        results, full_text, title, main_lang = check_article(doc, l)
        df       = pd.DataFrame(results)
        img_rows, tables_rows = analyse_figures_and_tables(doc, full_text, l)

    passed = sum(1 for r in results if r["Статус"] == "✅")
    warned = sum(1 for r in results if r["Статус"] == "⚠️")
    failed = sum(1 for r in results if r["Статус"] == "❌")
    total  = len(results)
    score  = int(passed / total * 100) if total > 0 else 0

    st.markdown(f"## {l['res_title']}")
    c1,c2,c3,c4,c5 = st.columns(5)
    c1.metric(l["total"],  total);  c2.metric(l["passed"], passed)
    c3.metric(l["warned"], warned); c4.metric(l["failed"], failed)
    c5.metric(l["score"],  f"{score}%")

    bar_clr = "#238636" if score >= 80 else "#d29922" if score >= 60 else "#da3633"
    bg_bar  = "#0f2340" if st.session_state.theme == "dark" else "#e9ecef"
    txt_clr = "#e2edf7" if st.session_state.theme == "dark" else "#ffffff"
    st.markdown(
        f'<div style="background:{bg_bar};border:1px solid #30363d;border-radius:6px;'
        f'height:28px;margin:8px 0 20px 0;">'
        f'<div style="background:{bar_clr};width:{score}%;height:28px;border-radius:6px;'
        f'display:flex;align-items:center;justify-content:center;'
        f'color:{txt_clr};font-weight:600;font-size:13px;">{score}%</div></div>',
        unsafe_allow_html=True)

    def highlight(row):  return [_ST.get(row["Статус"], _BASE)] * len(row)

    st.markdown(l["det_report"])
    st.dataframe(
        df.style.apply(highlight, axis=1),
        use_container_width=True, height=880,
        column_config={"№": st.column_config.NumberColumn(width="small")},
    )

    if img_rows:
        st.markdown(l["img_report"])
        df_img = pd.DataFrame(img_rows)
        scol   = l["img_status"]
        def hl_img(row): return [_ST.get(row[scol], _BASE)] * len(row)
        st.dataframe(
            df_img.style.apply(hl_img, axis=1),
            use_container_width=True,
        )

    if tables_rows:
        st.markdown("### 📊 Таблицы / Tables")
        df_tbl = pd.DataFrame(tables_rows)
        st.dataframe(df_tbl, use_container_width=True)

    st.markdown("---")
    ca, cb, cc = st.columns(3)
    bn = f"compliance_{title_for_filename(title)}"

    ca.download_button(
        l["btn_csv"],
        df.to_csv(index=False).encode("utf-8-sig"),
        f"{bn}.csv",
        "text/csv",
    )

    xb = BytesIO()
    with pd.ExcelWriter(xb, engine="openpyxl") as w:
        df.to_excel(w, index=False, sheet_name="Report")
        if img_rows:
            pd.DataFrame(img_rows).to_excel(w, index=False, sheet_name="Images")
        if tables_rows:
            pd.DataFrame(tables_rows).to_excel(w, index=False, sheet_name="Tables")
    cb.download_button(
        l["btn_xls"],
        xb.getvalue(),
        f"{bn}.xlsx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    if img_rows:
        st.download_button(
            l["btn_csv_fig"],
            pd.DataFrame(img_rows).to_csv(index=False).encode("utf-8-sig"),
            f"{bn}_figures.csv",
            "text/csv",
        )
    if tables_rows:
        st.download_button(
            l["btn_csv_tbl"],
            pd.DataFrame(tables_rows).to_csv(index=False).encode("utf-8-sig"),
            f"{bn}_tables.csv",
            "text/csv",
        )

    cc.download_button(
        l["btn_docx"],
        build_docx_report(results, l, total, passed, warned, failed, score,
                          img_rows, tables_rows),
        f"{bn}.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    problems = [r for r in results if r["Статус"] in ("❌", "⚠️")]
    if problems:
        st.markdown(l["req_fix"])
        for prob in problems:
            icon = "🔴" if prob["Статус"] == "❌" else "🟡"
            st.markdown(
                f"{icon} **{prob['Критерий']}** — {prob['Найдено']} "
                f"*({l['req']}: {prob['Требование']})*")
else:
    st.info(l["no_file"])

fc  = "#7b96b8" if st.session_state.theme == "dark" else "#555"
flk = "#58a6ff"  if st.session_state.theme == "dark" else "#0969da"
st.markdown("---")
st.markdown(
    f'<div style="text-align:center;font-size:12px;color:{fc};padding:12px 0 20px 0;line-height:2.2;">'
    f'<b style="font-size:13px;">© 2025 {l["f_author"]}</b><br>'
    f'📧 <a href="mailto:samarkhanov_kb@enu.kz" style="color:{flk};text-decoration:none;">samarkhanov_kb@enu.kz</a>'
    f'&nbsp;·&nbsp;'
    f'<a href="mailto:kanat.baurzhanuly@gmail.com" style="color:{flk};text-decoration:none;">kanat.baurzhanuly@gmail.com</a><br>'
    f'🏛️ <a href="https://fns.enu.kz/kz/page/departments/physical-and-economical-geography/faculty-members"'
    f'     target="_blank" style="color:{flk};text-decoration:none;">{l["f_univ"]}</a><br>'
    f'📄 {l["f_license"]}:&nbsp;'
    f'<a href="https://creativecommons.org/licenses/by/4.0/" target="_blank" style="color:{flk};text-decoration:none;">'
    f'CC BY 4.0 — Creative Commons Attribution 4.0 International</a>'
    f'</div>',
    unsafe_allow_html=True)
